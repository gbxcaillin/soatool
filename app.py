"""
SOA Completion Agent — Python Backend
Brightday Australia
Version: 1.2

Reads:  Fact Finder (.xlsx)  →  Fact Finder tab (via python-calamine)
        SOA Template (.docx) →  find & replace {{codes}} (via python-docx)

Outputs: Completed SOA draft (.docx) with all insertions in red font.
         Unmapped codes are left as raw {{code}} text.

Dependencies:
    pip install flask python-docx python-calamine gunicorn gevent

Run locally:
    python app.py
    Open http://localhost:5000

Deploy to Render:
    Start Command: gunicorn app:app --bind 0.0.0.0:$PORT --workers 1 --timeout 120
    Environment vars: SECRET_KEY, USERS (see DEPLOYMENT_GUIDE.md)
"""

from flask import Flask, request, send_file, jsonify, session, redirect, url_for, render_template_string
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date, timedelta
import io
import zipfile
import traceback
import re
import copy
import os
import hashlib

app = Flask(__name__)

# ─────────────────────────────────────────────
# AUTH CONFIG
# Read credentials from environment variables.
# Set these in Render dashboard — never hardcode.
# ─────────────────────────────────────────────
app.secret_key = os.environ.get("SECRET_KEY", "change-this-in-production")

# USERS dict — username: hashed password
# To generate a hash: python3 -c "import hashlib; print(hashlib.sha256('yourpassword'.encode()).hexdigest())"
# Add as many users as needed in the USERS env var format:
#   USERS=username1:hash1,username2:hash2
def load_users():
    users_env = os.environ.get("USERS", "")
    users = {}
    for entry in users_env.split(","):
        entry = entry.strip()
        if ":" in entry:
            username, pw_hash = entry.split(":", 1)
            users[username.strip().lower()] = pw_hash.strip()
    return users

def check_password(username, password):
    users = load_users()
    pw_hash = hashlib.sha256(password.encode()).hexdigest()
    return users.get(username.lower()) == pw_hash

def logged_in():
    return session.get("authenticated") is True

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
RED   = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Fact Finder tab — column letters for multi-fund fields (up to 5 funds)
# Columns: B=2, D=4, F=6, H=8, J=10
FUND_COLS = [2, 4, 6, 8, 10]


# ─────────────────────────────────────────────
# KYC FILE NOTE READER
# ─────────────────────────────────────────────

def read_kyc_note(docx_bytes):
    """
    Parse a KYC File Note (.docx) and return structured data the SOA agent can use:
        - risk_profile          : str  (from Paraplanning Request → Risk Profile)
        - scope                 : dict (super/insurance/salary_sacrifice/estate_planning -> 'in' | 'out')
        - goals                 : dict (5 verbatim goal paragraphs keyed by section)
        - meta                  : dict (client_name, adviser, meeting_date, platform, model)

    Relies on the standard KYC heading structure:
        Heading 1: 'Paraplanning Request', 'Client Goals Summary', 'Risk Profile Questions...'
        Heading 2: 'Superannuation Goals — Scoped in/out [limited to ...]'
                   'Insurance Goals — Scoped in/out [limited to ...]'
                   'Super Contribution Goals — Scoped in/out [limited to ...]'
                   'Estate Planning Goals — Scoped in/out [limited to ...]'
                   'Future Considerations — Retirement Goal'
    """
    from docx import Document as _DocxDocument
    doc = _DocxDocument(io.BytesIO(docx_bytes))

    # Walk paragraphs once, building a flat list of (level, heading, body_text) sections.
    # level 0 = pre-heading text, level 1/2 = real headings.
    sections = []
    current = {"level": 0, "heading": "", "text": ""}
    sections.append(current)

    def _norm_dashes(s):
        # Normalise unicode dashes so 'Scoped in' detection works regardless of em/en/hyphen
        return s.replace("—", "-").replace("–", "-").replace("−", "-")

    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = (p.style.name if p.style else "").strip()
        is_h1 = style_name.startswith("Heading 1")
        is_h2 = style_name.startswith("Heading 2")
        if is_h1 or is_h2:
            current = {"level": 1 if is_h1 else 2, "heading": text, "text": ""}
            sections.append(current)
        elif text:
            if current["text"]:
                current["text"] += "\n\n"
            current["text"] += text

    def find_section(*keywords, level=None):
        """First section whose heading contains all keywords (case-insensitive)."""
        for s in sections:
            h = s["heading"].lower()
            if all(k.lower() in h for k in keywords) and (level is None or s["level"] == level):
                return s
        return None

    # ── Paraplanning Request: extract Platform, Risk Profile, Model ──
    # Try the heading-anchored section first, but fall back to scanning the full
    # document text — some KYC notes use plain bold rather than Heading 1 for
    # 'Paraplanning Request'. The 'Platform:'/'Risk Profile:'/'Model:' labels are
    # distinctive enough that a full-doc scan won't false-match.
    risk_profile = ""
    platform     = ""
    model        = ""
    full_doc_text = "\n".join(p.text for p in doc.paragraphs)
    paraplanning = find_section("paraplanning")
    search_body = paraplanning["text"] if paraplanning else full_doc_text
    m = re.search(r"^\s*Platform:\s*(.+?)\s*$",       search_body, re.IGNORECASE | re.MULTILINE)
    if m: platform = m.group(1).strip()
    m = re.search(r"^\s*Risk Profile:\s*(.+?)\s*$",   search_body, re.IGNORECASE | re.MULTILINE)
    if m: risk_profile = m.group(1).strip()
    m = re.search(r"^\s*Model:\s*(.+?)\s*$",          search_body, re.IGNORECASE | re.MULTILINE)
    if m: model = m.group(1).strip()
    # Final fallback: if the section-scoped search didn't find them, sweep the whole doc
    if paraplanning and not (risk_profile and platform and model):
        if not platform:
            m = re.search(r"^\s*Platform:\s*(.+?)\s*$", full_doc_text, re.IGNORECASE | re.MULTILINE)
            if m: platform = m.group(1).strip()
        if not risk_profile:
            m = re.search(r"^\s*Risk Profile:\s*(.+?)\s*$", full_doc_text, re.IGNORECASE | re.MULTILINE)
            if m: risk_profile = m.group(1).strip()
        if not model:
            m = re.search(r"^\s*Model:\s*(.+?)\s*$", full_doc_text, re.IGNORECASE | re.MULTILINE)
            if m: model = m.group(1).strip()

    # ── Goal sections + scope tags ──
    # Heading examples: 'Superannuation Goals – Scoped in', 'Super Contribution Goals – Scoped out',
    # 'Estate Planning Goals – Scoped in limited to beneficiaries'.
    GOAL_KEYS = [
        # (key,                heading_keywords_to_match)
        ("super",              ["superannuation goals"]),
        ("insurance",          ["insurance goals"]),
        ("salary_sacrifice",   ["super contribution goals"]),
        ("estate_planning",    ["estate planning goals"]),
        ("retirement",         ["retirement goal"]),  # under 'Future Considerations'
    ]

    goals = {}
    scope = {}
    for key, keywords in GOAL_KEYS:
        sect = None
        for kw in keywords:
            sect = find_section(kw, level=2)
            if sect: break
        if sect:
            goals[key] = sect["text"].strip()
            heading_norm = _norm_dashes(sect["heading"]).lower()
            if "scoped out" in heading_norm:
                scope[key] = "out"
            elif "scoped in" in heading_norm:
                scope[key] = "in"
            else:
                scope[key] = "in"   # retirement section has no scope tag; treat as 'in'
        else:
            goals[key] = ""
            scope[key] = "in"

    # ── File Note metadata (client, adviser, date) — informational only for now ──
    client_name  = ""
    adviser      = ""
    meeting_date = ""
    m = re.search(r"^\s*Client:\s*(.+?)\s*$",  full_doc_text, re.IGNORECASE | re.MULTILINE)
    if m: client_name = m.group(1).strip()
    m = re.search(r"^\s*Adviser:\s*(.+?)\s*$", full_doc_text, re.IGNORECASE | re.MULTILINE)
    if m: adviser = m.group(1).strip()
    m = re.search(r"^\s*Date:\s*(.+?)\s*$",    full_doc_text, re.IGNORECASE | re.MULTILINE)
    if m: meeting_date = m.group(1).strip()

    return {
        "risk_profile": risk_profile,
        "scope":        scope,
        "goals":        goals,
        "meta": {
            "client_name":  client_name,
            "adviser":      adviser,
            "meeting_date": meeting_date,
            "platform":     platform,
            "model":        model,
        },
    }


# ─────────────────────────────────────────────
# SCENARIO TRAINING DOCUMENT (LIBRARY) READER
# ─────────────────────────────────────────────

# Markers we recognise inside the training document and the SOA template.
# Matches: {{ScenarioNa}} … {{ScenarioNi}}, {{ScenarioNoptIn}}, {{ScenarioNAdviceLimitation}}.
SCENARIO_MARKER_RE = re.compile(
    r'\{\{Scenario([1-7])(?:([a-z])|(optIn|AdviceLimitation))\}\}'
)

def _get_para_text(el):
    """Concatenated text of all <w:t> children inside a body element (paragraph or table)."""
    return "".join(t.text or "" for t in el.iter(qn('w:t')))


def read_scenario_library(docx_bytes):
    """
    Parse the Scenario Training .docx into a dict of marker -> [body XML elements].

    Each scenario marker in the training doc is followed by the content that should be
    spliced into the SOA template at that marker location. Content runs until the next
    scenario marker (or end of doc). Content includes paragraphs, tables, lists,
    formatting, and any nested {{codes}} (which get resolved by the later find-and-replace
    pass once spliced into the SOA template).

    Returns: dict { '{{Scenario1a}}': [<w:p>, <w:tbl>, ...], ... }
    """
    from docx import Document as _DocxDocument
    doc = _DocxDocument(io.BytesIO(docx_bytes))
    body = doc.element.body

    library = {}
    current_marker = None
    current_elements = []

    # Matches non-marker section headings like "Scenario - (1) - Has group insurance...".
    # These divide the training doc into sections but aren't content to insert.
    SECTION_DIVIDER_RE = re.compile(r'^\s*Scenario\s*-\s*\(\s*[1-7]\s*\)\s*-', re.IGNORECASE)

    for el in list(body):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'sectPr':
            continue
        text = _get_para_text(el) if tag in ('p', 'tbl') else ""
        # Section divider — close out the current marker; the divider itself is not content
        if tag == 'p' and SECTION_DIVIDER_RE.match(text):
            if current_marker:
                library[current_marker] = current_elements
            current_marker = None
            current_elements = []
            continue
        m = SCENARIO_MARKER_RE.search(text) if tag == 'p' else None
        if m:
            if current_marker:
                # If the same marker appears twice (template typo), last write wins
                # — flag in logs but proceed.
                if m.group() in library:
                    # Restore previous; will be overwritten below
                    pass
                library[current_marker] = current_elements
            current_marker = m.group()
            current_elements = []
            continue
        if current_marker:
            current_elements.append(el)

    if current_marker:
        library[current_marker] = current_elements

    return library


def _load_bundled_scenario_library():
    """
    Load the bundled scenario_library.docx that ships alongside app.py — used as the
    default library when the adviser doesn't upload one via the UI. Returns None if
    the file is absent or fails to parse (the agent will then operate with no library
    and leave the selected scenario's markers raw).
    """
    try:
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "scenario_library.docx")
        if os.path.exists(path):
            with open(path, "rb") as f:
                return read_scenario_library(f.read())
    except Exception:
        pass
    return None


# Loaded once at process startup. Re-deploy to pick up a new scenario_library.docx.
DEFAULT_SCENARIO_LIBRARY = _load_bundled_scenario_library()


# ─────────────────────────────────────────────
# FACT FINDER READER
# ─────────────────────────────────────────────

def read_fact_finder(xlsx_bytes, risk_profile, no_insurance_flag,
                     no_trauma_flag=False, no_salsac_flag=False,
                     insurance_only_flag=False, scenario="",
                     goal_overrides=None, rollover_type=""):
    """
    Read the Fact Finder xlsx and return:
        - data dict  { "{{CODE}}": "value" }
        - conditionals dict  { "DELETE_KEY": True/False }

    UI flags:
        no_insurance_flag    -> drives {{DeleteIfNoInsuranceAtAll}} and {{DeleteIfNoInsuranceAdvice}}
        no_trauma_flag       -> drives {{DeleteIfNoScopedTrauma}}
        no_salsac_flag       -> drives {{DeleteIfNoSalarySacrificeAdvice}}
        insurance_only_flag  -> drives {{DeleteIfInsuranceOnlyClient}}
        scenario             -> "1".."6"; keeps {{ScenarioN}} block, deletes the other five
        goal_overrides       -> optional dict of goal text from the KYC note, keyed by:
                                'super', 'insurance', 'salary_sacrifice', 'estate_planning', 'retirement'.
                                When a key has a non-empty value, the corresponding {{*Goal}} placeholder
                                is replaced in the SOA. Empty / missing keys leave the placeholder raw.
    """
    from python_calamine import CalamineWorkbook
    cal_wb   = CalamineWorkbook.from_filelike(io.BytesIO(xlsx_bytes))
    sheet    = cal_wb.get_sheet_by_name("Fact Finder")
    raw_rows = sheet.to_python(skip_empty_area=False)

    # Build a row/col lookup dict identical to openpyxl's ws.cell(row, col).value
    # calamine returns a list of rows; each row is a list of values
    # rows and cols are 1-indexed to match existing code
    cell_data = {}
    for r_idx, row in enumerate(raw_rows, start=1):
        for c_idx, val in enumerate(row, start=1):
            cell_data[(r_idx, c_idx)] = val

    def cell(row, col):
        """Return cleaned string value from a cell, or '' if empty/zero placeholder.
        Per spec: a numeric 0 (int or float) is treated as empty."""
        v = cell_data.get((row, col))
        if v is None:
            return ""
        # Numeric zero -> empty (covers calamine returning 0.0 for blank-looking cells)
        if isinstance(v, (int, float)) and not isinstance(v, bool) and float(v) == 0:
            return ""
        s = str(v).strip()
        if s in ("", "0", "0.0", "00:00:00", "#REF!", "None"):
            return ""
        return s

    def cells_across(row, cols=FUND_COLS):
        """Return list of non-empty values across multiple fund columns."""
        return [cell(row, c) for c in cols if cell(row, c)]

    def join_funds(row, sep=", "):
        vals = cells_across(row)
        return sep.join(vals) if vals else ""

    def sum_funds(row):
        total = 0
        for c in FUND_COLS:
            v = cell_data.get((row, c))
            try:
                total += float(str(v).replace(",", "").replace("$", ""))
            except Exception:
                pass
        return total

    def currency(row, col):
        """Currency-format a cell value. Returns '' for empty cells AND numeric 0."""
        v = cell_data.get((row, col))
        if v is None:
            return ""
        try:
            n = float(str(v).replace(',', '').replace('$', ''))
        except (ValueError, TypeError):
            return ""
        if n == 0:
            return ""
        return f"${n:,.0f}"

    def currency_sum(row):
        s = sum_funds(row)
        return f"${s:,.0f}" if s else ""

    def age_from_dob(row, col=2):
        v = cell_data.get((row, col))
        if not v:
            return None
        try:
            from datetime import datetime
            if hasattr(v, 'year'):
                dob = v
            else:
                for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
                    try:
                        dob = datetime.strptime(str(v), fmt)
                        break
                    except ValueError:
                        continue
                else:
                    return None
            today = date.today()
            return today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
        except Exception:
            return None

    def _whole_num(s):
        """Strip a trailing '.0' from a value that's an integer in disguise.
        Calamine returns numeric cells as floats (e.g. postcode 3137 -> 3137.0).
        Returns input unchanged if it isn't a clean integer."""
        if not s:
            return s
        try:
            n = float(s)
            return str(int(n)) if n == int(n) else str(s)
        except (ValueError, TypeError):
            return str(s)

    def format_au_phone(raw):
        """Format an Australian phone number as '0XXX XXX XXX'.
        Handles inputs with spaces, dashes, +61 prefix, or already-formatted.
        Returns the original string if it doesn't look like a 10-digit AU number."""
        if not raw:
            return ""
        digits = "".join(c for c in str(raw) if c.isdigit())
        # Convert +61... or 61... (11 digits) to 0...
        if len(digits) == 11 and digits.startswith("61"):
            digits = "0" + digits[2:]
        if len(digits) == 10 and digits.startswith("0"):
            return f"{digits[:4]} {digits[4:7]} {digits[7:]}"
        return str(raw).strip()

    def format_date(row, col=2):
        v = cell_data.get((row, col))
        if not v:
            return ""
        try:
            if hasattr(v, 'strftime'):
                return v.strftime("%d/%m/%Y")
            return str(v)
        except Exception:
            return str(v)

    # ── Personal Details ──
    title       = cell(10, 2)
    first_name  = cell(11, 2)
    last_name   = cell(13, 2)
    # Per spec: ClientFullName = first + last only (middle name skipped)
    full_name_parts = [p for p in [first_name, last_name] if p]
    full_name   = " ".join(full_name_parts)
    dob_str     = format_date(15, 2)
    age         = age_from_dob(15, 2)
    phone       = format_au_phone(cell(16, 2))
    email       = cell(17, 2)
    # Strip trailing '.0' on each address part so the postcode (B21) doesn't render as '3137.0'.
    address_parts = [_whole_num(p) for p in [cell(18,2), cell(19,2), cell(20,2), cell(21,2)] if p]
    address     = ", ".join(address_parts)

    # ── Employment ──
    occupation  = cell(28, 2)
    emp_status  = cell(23, 2)

    # ── Income ──
    gross_income_raw = cell_data.get((32, 2))
    try:
        gross_income_num = float(str(gross_income_raw).replace(",","").replace("$",""))
        gross_income = f"${gross_income_num:,.0f}"
    except Exception:
        gross_income_num = 0
        gross_income = ""

    sgc_pct_raw = cell_data.get((34, 2))
    try:
        sgc_pct = float(str(sgc_pct_raw).replace("%","")) / 100
    except Exception:
        sgc_pct = 0.12
    super_contribution = f"${gross_income_num * sgc_pct:,.0f}" if gross_income_num else ""

    salary_sacrifice_raw = cell_data.get((35, 2))
    try:
        salary_sacrifice = f"${float(str(salary_sacrifice_raw).replace(',','').replace('$','')):,.0f}"
    except Exception:
        salary_sacrifice = ""

    # Personal deductible contributions — FF B36 (separate from salary sacrifice in B35).
    # Renamed from {{AnnualisedSalarySacrificeAmount}} (which read B35) to {{PersonalDeductibleContributions}}.
    personal_deductible_contributions_raw = cell_data.get((36, 2))
    try:
        personal_deductible_contributions = f"${float(str(personal_deductible_contributions_raw).replace(',','').replace('$','')):,.0f}"
    except Exception:
        personal_deductible_contributions = ""

    # Non-concessional contributions — FF B37. New for the v2.2 template.
    non_concessional_contributions_raw = cell_data.get((37, 2))
    try:
        non_concessional_contributions = f"${float(str(non_concessional_contributions_raw).replace(',','').replace('$','')):,.0f}"
    except Exception:
        non_concessional_contributions = ""

    # ── Retirement Age ──
    # Per spec: default to B8; use B9 only if B9 is non-empty AND differs from B8.
    # cell() returns "" for both blank and numeric 0, so the simple check covers both.
    # Then strip any trailing '.0' (calamine returns numeric ages as floats).
    ret_age_1 = cell(8, 2)
    ret_age_2 = cell(9, 2)
    retirement_age = ret_age_2 if (ret_age_2 and ret_age_2 != ret_age_1) else ret_age_1
    retirement_age = _whole_num(retirement_age) if retirement_age else retirement_age

    # ── Spouse ──
    spouse_dob  = format_date(47, 2)
    spouse_income_raw = cell_data.get((49, 2))
    try:
        spouse_income = f"${float(str(spouse_income_raw).replace(',','').replace('$','')):,.0f}"
    except Exception:
        spouse_income = ""
    spouse_balance_raw = cell_data.get((50, 2))
    try:
        spouse_balance = f"${float(str(spouse_balance_raw).replace(',','').replace('$','')):,.0f}"
    except Exception:
        spouse_balance = ""

    # ── Dependants ──
    has_spouse = bool(cell(46, 2))  # Spouse Name row
    dep_ages = cells_across(56)
    no_dependants = (1 if has_spouse else 0) + len(dep_ages)

    # ── Assets & Liabilities ──
    primary_residence_val = currency(73, 2)
    primary_residence_debt = currency(74, 2)
    investment_prop_val = currency(76, 2)
    investment_prop_debt = currency(77, 2)
    other_asset1_val = currency(79, 2)
    personal_loan1_val = currency(81, 2)

    # Helper: add cell value as float to a running total, skipping bad/empty values.
    def _add_cell_to_total(running_total, row, col):
        v = cell_data.get((row, col))
        try:
            return running_total + float(str(v).replace(",", "").replace("$", ""))
        except Exception:
            return running_total

    # Total assets — per v2.2 Code Map:
    #   B71  (cash / liquid assets)
    #   B73  (primary residence)
    #   row 76 across fund cols B/D/F/H/J  (investment property values)
    #   row 79 across fund cols B/D/F/H/J  (other asset values)
    # NOTE: super balances (sum of row 94) are NOT included in TotalAssetValue.
    total_assets = 0
    total_assets = _add_cell_to_total(total_assets, 71, 2)
    total_assets = _add_cell_to_total(total_assets, 73, 2)
    for c in FUND_COLS:
        total_assets = _add_cell_to_total(total_assets, 76, c)
        total_assets = _add_cell_to_total(total_assets, 79, c)
    total_assets_str = f"${total_assets:,.0f}" if total_assets else ""

    # Total liabilities — per v2.2 Code Map:
    #   B74  (primary residence debt)
    #   row 77 across fund cols B/D/F/H/J  (investment property debts)
    #   row 81 across fund cols B/D/F/H/J  (personal loans)
    total_liabilities = 0
    total_liabilities = _add_cell_to_total(total_liabilities, 74, 2)
    for c in FUND_COLS:
        total_liabilities = _add_cell_to_total(total_liabilities, 77, c)
        total_liabilities = _add_cell_to_total(total_liabilities, 81, c)
    total_liabilities_str = f"${total_liabilities:,.0f}" if total_liabilities else ""

    # ── Super Funds ──
    current_super_funds = join_funds(92)
    current_super_balance = currency_sum(94)
    current_balance = current_super_balance

    # ── Insurance across funds ──
    def insurance_across(row, currency=False):
        """Join non-empty values from row across fund columns with ' / '.
        - Numeric 0 / blank cells are skipped.
        - Trailing '.0' is stripped from whole numbers (calamine returns floats for numeric cells).
        - When currency=True, numeric values are formatted as $X,XXX (no decimals).
        - Non-numeric strings (e.g. 'Yes', '$5/wk') are kept as-is.
        """
        vals = []
        for c in FUND_COLS:
            v = cell_data.get((row, c))
            if v is None:
                continue
            # Numeric zero -> treat as empty
            if isinstance(v, (int, float)) and not isinstance(v, bool) and float(v) == 0:
                continue
            s = str(v).strip()
            if s in ("", "0", "0.0", "None"):
                continue
            # Try to interpret as a number (handles "400000.0", "400,000", "$400000")
            try:
                n = float(s.replace(",", "").replace("$", ""))
                if currency:
                    s_clean = f"${n:,.0f}"
                elif n == int(n):
                    s_clean = str(int(n))
                else:
                    s_clean = s
            except (ValueError, TypeError):
                s_clean = s
            vals.append(s_clean)
        if not vals:
            return ""
        unique = list(dict.fromkeys(vals))
        return " / ".join(unique)

    life_ins   = insurance_across(102, currency=True)   # Life cover $
    tpd_ins    = insurance_across(103, currency=True)   # TPD cover $
    ip_month   = insurance_across(104, currency=True)   # IP monthly benefit $
    ip_wait    = insurance_across(105)                  # waiting period (duration)
    ip_benefit = insurance_across(106)                  # benefit period (duration)
    premiums   = insurance_across(107, currency=True)   # premiums $

    # ── Binding Death Nominee ──
    # Per spec: read from row 63 across fund columns (was row 62).
    # cell() now treats numeric 0 / blank as empty so zero-cells are skipped.
    # Joining: 'N/A' if none, single name as-is, 'A and B' if exactly 2,
    # comma-separated for 3+.
    nominee_names = [cell(63, c) for c in FUND_COLS if cell(63, c)]
    if not nominee_names:
        binding_death_nominee = "N/A"
    elif len(nominee_names) == 1:
        binding_death_nominee = nominee_names[0]
    elif len(nominee_names) == 2:
        binding_death_nominee = " and ".join(nominee_names)
    else:
        binding_death_nominee = ", ".join(nominee_names)

    # ── Current Date ──
    current_date = date.today().strftime("%d %B %Y")

    # ── Ongoing Fee Agreement (OFA) dates ──
    # Presentation date = today + 2 days (when the OFA is presented to the client).
    # Reference date    = presentation date + 12 months (renewal review).
    # Arrangement end   = reference date + 150 days (consent expiry, per the OFA template body).
    # The placeholder text in the OFA template says '+ 5 months' for the second date,
    # but the body and the spec confirm 150 days — using 150 days here.
    import calendar as _cal
    def _add_months(d, months):
        new_month_idx = d.month - 1 + months
        new_year = d.year + new_month_idx // 12
        new_month = new_month_idx % 12 + 1
        new_day = min(d.day, _cal.monthrange(new_year, new_month)[1])
        return date(new_year, new_month, new_day)

    _today = date.today()
    presentation_date    = _today + timedelta(days=2)
    reference_date       = _add_months(presentation_date, 12)
    arrangement_end_date = reference_date + timedelta(days=150)
    DATE_FMT_LONG = "%d %B %Y"
    reference_date_str       = reference_date.strftime(DATE_FMT_LONG)
    arrangement_end_date_str = arrangement_end_date.strftime(DATE_FMT_LONG)

    # ── Ongoing Advice Fee ({{zzz}} = annual, {{zzzMonthly}} = approx monthly) ──
    # 1.1% of total super balance, capped at $10,000 p.a.
    # Output is the bare number "X,XXX.XX" — the OFA template already has '$' before the placeholder.
    # Monthly is the annual figure divided by 12 (approx — actual monthly deduction may vary).
    _total_super = sum_funds(94)
    _ongoing_fee = min(_total_super * 0.011, 10_000.0)
    _ongoing_fee_monthly = _ongoing_fee / 12 if _ongoing_fee > 0 else 0
    ongoing_fee_str         = f"{_ongoing_fee:,.2f}"          if _ongoing_fee > 0 else ""
    ongoing_fee_monthly_str = f"{_ongoing_fee_monthly:,.2f}"  if _ongoing_fee_monthly > 0 else ""

    # ── Upfront Advice Fee — %ofSuperBalance breakdown (Pearl X v2.2 template) ──
    # Total upfront fee = min($4,950, super balance × 3.3%). Below $150k it scales;
    # above that it caps at $4,950.
    # The percentage placeholders are sub-portions of that total fee:
    #   {{100%ofSuperBalance}} = total upfront fee
    #   {{45%ofSuperBalance}}  = 45% of total  (Advice preparation fee)
    #   {{22.5%ofSuperBalance}}= 22.5% of total (Strategic AND Implementation — same value used twice)
    #   {{10%ofSuperBalance}}  = 10% of total  (Licensee fee)
    # Output bare numbers "X,XXX.XX" — template has '$' before the placeholder.
    _upfront_fee_total = min(4_950.0, _total_super * 0.033) if _total_super > 0 else 0
    if _upfront_fee_total > 0:
        upfront_fee_100_str  = f"{_upfront_fee_total:,.2f}"
        upfront_fee_45_str   = f"{_upfront_fee_total * 0.45:,.2f}"
        upfront_fee_225_str  = f"{_upfront_fee_total * 0.225:,.2f}"
        upfront_fee_10_str   = f"{_upfront_fee_total * 0.10:,.2f}"
    else:
        upfront_fee_100_str = upfront_fee_45_str = upfront_fee_225_str = upfront_fee_10_str = ""

    # ── Risk Profile (from UI selection) ──
    current_risk_profile = risk_profile  # passed in from form

    # ─────────────────────────────────
    # BUILD DATA DICT
    # ─────────────────────────────────
    data = {
        "{{Title}}":                             title,
        "{{ClientFullName}}":                    full_name,
        "{{ClientFirstName}}":                   first_name,
        "{{ClientLastName}}":                    last_name,
        "{{ClientDOB}}":                         dob_str,
        "{{ClientAddress}}":                     address,
        "{{ClientPhone}}":                       phone,
        "{{ClientEmail}}":                       email,
        "{{ClientOccupation}}":                  occupation,
        "{{ClientSalary}}":                      gross_income,
        "{{fld_SuperContribution}}":             super_contribution,
        "{{fld_SalarySacrifice}}":               salary_sacrifice,
        "{{CurrentSuperFunds}}":                 current_super_funds,
        "{{SpouseDOB}}":                         spouse_dob,
        "{{SpouseIncome}}":                      spouse_income,
        "{{SpouseBalance}}":                     spouse_balance,
        "{{NoDependants}}":                      str(no_dependants),
        "{{fld_CurrentSuperannuationBalance}}":  current_super_balance,
        "{{CurrentLifeInsurance}}":              life_ins,
        "{{CurrentTPDInsurance}}":               tpd_ins,
        "{{CurrentIncomeProtectionPerMonth}}":   ip_month,
        "{{CurrentIncomeProtectionWaitingPeriod}}": ip_wait,
        "{{CurrentIncomeProtectionBenefitPeriod}}": ip_benefit,
        "{{CurrentSuperPremiums}}":              premiums,
        "{{ValueOfPrimaryResidence}}":           primary_residence_val,
        "{{DebtOnPrimaryResidence}}":            primary_residence_debt,
        "{{ValueOfInvestmentProperty}}":         investment_prop_val,
        "{{DebtOnInvestmentProperty}}":          investment_prop_debt,
        "{{OtherAsset1Value}}":                  other_asset1_val,
        "{{PersonalLoan1Value}}":                personal_loan1_val,
        "{{TotalAssetValue}}":                   total_assets_str,
        "{{TotalLiabilityValue}}":               total_liabilities_str,
        "{{RetirementAge}}":                     retirement_age,
        "{{CurrentBalance}}":                    current_balance,
        "{{CurrentAge}}":                        str(age) if age else "",
        "{{CurrentDate}}":                       current_date,
        # Ongoing Fee Agreement date placeholders — exact strings (with leading spaces) per the OFA template.
        "{{ Presentation date + 12 months}}":    reference_date_str,
        "{{ Reference date + 5 months}}":        arrangement_end_date_str,
        # OFA ongoing advice fee — 1.1% of total super balance, capped at $10,000 p.a.
        "{{zzz}}":                               ongoing_fee_str,           # annual
        "{{zzzMonthly}}":                        ongoing_fee_monthly_str,   # approx monthly = annual / 12
        # Upfront advice fee breakdown — total = min($4,950, balance × 3.3%). See helper above.
        "{{100%ofSuperBalance}}":                upfront_fee_100_str,       # total upfront fee
        "{{45%ofSuperBalance}}":                 upfront_fee_45_str,        # advice preparation
        "{{22.5%ofSuperBalance}}":               upfront_fee_225_str,       # strategic AND implementation
        "{{10%ofSuperBalance}}":                 upfront_fee_10_str,        # licensee
        "{{PersonalDeductibleContributions}}":   personal_deductible_contributions,   # FF B36 (was {{AnnualisedSalarySacrificeAmount}} reading B35)
        "{{NonConcessionalContributions}}":      non_concessional_contributions,      # FF B37 — new for v2.2
        "{{BindingDeathNominee}}":               binding_death_nominee,
        "{{CurrentRiskProfile}}":                current_risk_profile,
        "{{EmploymentStatus}}":                  emp_status,
        "{{MaritalStatus}}":                     cell(45, 2),
        # Goals — left as raw codes (adviser-completed)
        # Table codes — left as raw codes (adviser-completed)
        # $r1-$r4, $p1-$p4 — adviser-completed (different placeholder syntax, regex ignores them)
    }

    # ─────────────────────────────────
    # CONDITIONALS
    # ─────────────────────────────────
    total_balance = sum_funds(94)
    super_contribution_num = gross_income_num * sgc_pct if gross_income_num else 0

    # Row 100: insurance in fund — check all fund columns
    has_any_insurance = any(
        str(cell_data.get((100, c)) or "").strip().lower() == "yes"
        for c in FUND_COLS
    )

    # Row 108: medically underwritten
    has_underwritten = any(
        str(cell_data.get((108, c)) or "").strip().lower() == "medically underwritten"
        for c in FUND_COLS
    )

    # Count of distinct super funds (non-empty values across row 92 fund-name columns).
    # Used by {{DeleteIfOneSuperFund}} — only one fund means there's nothing to consolidate.
    fund_count = sum(1 for c in FUND_COLS if cell(92, c))

    conditionals = {
        # True = DELETE this block
        "DeleteIfAgeGreaterThan55":              (age is not None and age >= 55),
        "DeleteIfAgeLessThan55":                 (age is not None and age < 55),
        "DeleteIfBalanceBelow450k":              (total_balance < 450_000),
        "DeleteIfSuperContributionsBelow27k":    (super_contribution_num < 27_000),
        "DeleteIfNoCurrentInsurance":            (not has_any_insurance),
        "DeleteIfNoInsuranceAtAll":              no_insurance_flag,        # legacy UI checkbox
        "DeleteIfNoInsuranceAdvice":             no_insurance_flag,        # new spec name, same checkbox
        "DeleteIfNoScopedTrauma":                no_trauma_flag,           # UI checkbox
        "DeleteIfNoSalarySacrificeAdvice":       no_salsac_flag,           # UI checkbox
        "DeleteIfInsuranceOnlyClient":           insurance_only_flag,      # UI checkbox
        # Per Code Map: delete the no-scoped-insurance block when adviser picks scenario 5 or 6
        "DeleteIfNoScopedInsurance":             (scenario in {"5", "6"}),
        "DeleteIfNoTrauma":                      no_trauma_flag,           # alias of NoScopedTrauma
        # Per Code Map v2: delete the personal-deductible-contributions block when both
        # B35 (salary sacrifice) AND B36 are EMPTY / zero. cell() treats 0 as empty,
        # so empty-string checks cover both blank and numeric-zero cells.
        "DeleteIfPersonalDeductibleContributions": (not cell(35, 2)) and (not cell(36, 2)),
        # Per Code Map: total liabilities = B74 + B77 + B81 (computed above as total_liabilities)
        "DeleteIfNoDebts":                       (total_liabilities == 0),
        "DeleteIfClientHasDebts":                (total_liabilities > 0),
        "DeleteifNoCurrentUnderwrittenInsurance": (not has_underwritten),
        # v2.2 new conditionals
        "DeleteIfOneSuperFund":                  (fund_count == 1),
        # Rollover dropdown (Step 2): "Full" / "Partial". Each conditional's name reads
        # as "delete THIS block IF doing the named rollover type".
        # The DeleteIfFullRollover marker wraps the PARTIAL rollover prose; we want to
        # delete the partial prose when the adviser picks Full → True when "Full".
        # The DeleteIfPartialRollover marker wraps the FULL rollover prose; we want to
        # delete the full prose when the adviser picks Partial → True when "Partial".
        "DeleteIfFullRollover":                  (rollover_type == "Full"),
        "DeleteIfPartialRollover":               (rollover_type == "Partial"),
    }

    # Scenarios 1–6: keep selected, delete the other five.
    # If no scenario selected, all six are kept (markers stripped).
    for n in range(1, 8):
        conditionals[f"DeleteScenario{n}"] = (scenario != "" and scenario != str(n))

    # ── Goal overrides from KYC note ──
    # Only inject goals that have non-empty text. Codes for empty goals stay raw
    # (matching legacy adviser-completed behavior).
    if goal_overrides:
        GOAL_CODE_MAP = {
            "super":            "{{SuperGoal}}",
            "insurance":        "{{InsuranceGoal}}",
            "salary_sacrifice": "{{SalarySacrificeGoal}}",
            "estate_planning":  "{{EstatePlanningGoal}}",
            "retirement":       "{{RetirementGoal}}",
        }
        for k, code in GOAL_CODE_MAP.items():
            v = (goal_overrides.get(k) or "").strip()
            if v:
                data[code] = v

    return data, conditionals


# ─────────────────────────────────────────────
# SOA DOCUMENT PROCESSOR
# ─────────────────────────────────────────────

# Codes that are intentionally left as raw {{code}} — never replaced
UNMAPPED_CODES = {
    "{{Date}}",
    "{{OtherAsset1}}",
    "{{OtherAsset2}}",
    "{{OtherAsset2Value}}",
    "{{PersonalLoan2Value}}",
    "{{PersonalLoan1}}",
    "{{PersonalLoan2}}",
    "{{NeedsAnalysisLifeInsurance}}",
    "{{NeedsAnalysisTPD}}",
    "{{NeedsAnalysisIP}}",
    "{{NeedsAnalysisTrauma}}",
    "{{Tbl_SalarySacrifice}}",
    "{{tbl_CurrentSuperFundsRiskProfilePerformance}}",
    "{{Make personal deductible contributions/Salary sacrifice}}",
    "{{CurrentInsuer}}",
    "{{SalarySacrificeAmount}}",
    "{{SalarySacrificeFrequency}}",
    "{{NetTaxSavings}}",
    "{{SuperGoal}}",
    "{{InsuranceGoal}}",
    "{{SalarySacrificeGoal}}",
    "{{EstatePlanningGoal}}",
    "{{RetirementGoal}}",
    "{{DeleteIfNoInsuranceAtAll}}",
    "{{EndDeleteIfNoInsuranceAtAll}}",
}

# Pair up conditional block tags
CONDITIONAL_PAIRS = [
    ("{{DeleteIfAgeGreaterThan55}}",              "{{EndDeleteIfAgeGreaterThan55}}",              "DeleteIfAgeGreaterThan55"),
    ("{{DeleteIfAgeLessThan55}}",                 "{{EndDeleteIfAgeLessThan55}}",                 "DeleteIfAgeLessThan55"),
    ("{{DeleteIfBalanceBelow450k}}",              "{{EndDeleteIfBalanceBelow450k}}",              "DeleteIfBalanceBelow450k"),
    ("{{DeleteIfSuperContributionsBelow27k}}",    "{{EndDeleteIfSuperContributionsBelow27k}}",    "DeleteIfSuperContributionsBelow27k"),
    ("{{DeleteIfNoCurrentInsurance}}",            "{{EndDeleteIfNoCurrentInsurance}}",            "DeleteIfNoCurrentInsurance"),
    ("{{DeleteifNoCurrentUnderwrittenInsurance}}","{{EndDeleteifNoCurrentUnderwrittenInsurance}}","DeleteifNoCurrentUnderwrittenInsurance"),
    # New spec — driven by UI checkboxes
    ("{{DeleteIfNoInsuranceAdvice}}",             "{{EndDeleteIfNoInsuranceAdvice}}",             "DeleteIfNoInsuranceAdvice"),
    ("{{DeleteIfNoScopedTrauma}}",                "{{EndDeleteIfNoScopedTrauma}}",                "DeleteIfNoScopedTrauma"),
    ("{{DeleteIfNoTrauma}}",                      "{{EndDeleteIfNoTrauma}}",                      "DeleteIfNoTrauma"),
    ("{{DeleteIfNoSalarySacrificeAdvice}}",       "{{EndDeleteIfNoSalarySacrificeAdvice}}",       "DeleteIfNoSalarySacrificeAdvice"),
    ("{{DeleteIfInsuranceOnlyClient}}",           "{{EndDeleteIfInsuranceOnlyClient}}",           "DeleteIfInsuranceOnlyClient"),
    # Auto from FF / scenario selection
    ("{{DeleteIfNoScopedInsurance}}",             "{{EndDeleteIfNoScopedInsurance}}",             "DeleteIfNoScopedInsurance"),
    ("{{DeleteIfNoDebts}}",                       "{{EndDeleteIfNoDebts}}",                       "DeleteIfNoDebts"),
    ("{{DeleteIfClientHasDebts}}",                "{{EndDeleteIfClientHasDebts}}",                "DeleteIfClientHasDebts"),
    ("{{DeleteIfPersonalDeductibleContributions}}","{{EndDeleteIfPersonalDeductibleContributions}}","DeleteIfPersonalDeductibleContributions"),
    # v2.2 new conditional pairs
    ("{{DeleteIfOneSuperFund}}",                  "{{EndDeleteIfOneSuperFund}}",                  "DeleteIfOneSuperFund"),
    ("{{DeleteIfFullRollover}}",                  "{{EndDeleteIfFullRollover}}",                  "DeleteIfFullRollover"),
    ("{{DeleteIfPartialRollover}}",               "{{EndDeleteIfPartialRollover}}",               "DeleteIfPartialRollover"),
    # Scenario 1–6 — adviser picks one in the UI; the other five are deleted
    ("{{Scenario1}}", "{{EndScenario1}}", "DeleteScenario1"),
    ("{{Scenario2}}", "{{EndScenario2}}", "DeleteScenario2"),
    ("{{Scenario3}}", "{{EndScenario3}}", "DeleteScenario3"),
    ("{{Scenario4}}", "{{EndScenario4}}", "DeleteScenario4"),
    ("{{Scenario5}}", "{{EndScenario5}}", "DeleteScenario5"),
    ("{{Scenario6}}", "{{EndScenario6}}", "DeleteScenario6"),
    ("{{Scenario7}}", "{{EndScenario7}}", "DeleteScenario7"),
]

# For DeleteIfNoInsuranceAtAll — single tag (no end tag), marks start of section to delete
# We treat the content following it until next section heading as the block
NO_INSURANCE_SINGLE_TAG = "{{DeleteIfNoInsuranceAtAll}}"


def get_full_text(paragraph):
    return "".join(run.text for run in paragraph.runs)


def para_contains(paragraph, code):
    return code in get_full_text(paragraph)


def replace_code_in_run(run, code, value, use_red):
    """Replace a code in a single run, applying red font to the replacement."""
    if code not in run.text:
        return
    parts = run.text.split(code)
    # If only one part before and after — simple case
    if len(parts) == 2:
        before, after = parts
        run.text = before
        # Insert red replacement run after this run
        p = run._r.getparent()
        idx = list(p).index(run._r)

        def make_run(text, red):
            from docx.oxml import OxmlElement
            r_el = OxmlElement('w:r')
            # Copy rPr from original run
            if run._r.find(qn('w:rPr')) is not None:
                rPr = copy.deepcopy(run._r.find(qn('w:rPr')))
                # Set or remove colour
                color_el = rPr.find(qn('w:color'))
                if color_el is None:
                    color_el = OxmlElement('w:color')
                    rPr.append(color_el)
                if red:
                    color_el.set(qn('w:val'), 'FF0000')
                else:
                    color_el.set(qn('w:val'), 'auto')
                r_el.append(rPr)
            else:
                if red:
                    rPr = OxmlElement('w:rPr')
                    color_el = OxmlElement('w:color')
                    color_el.set(qn('w:val'), 'FF0000')
                    rPr.append(color_el)
                    r_el.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.text = text
            if text.startswith(' ') or text.endswith(' '):
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_el.append(t_el)
            return r_el

        # Insert replacement
        if value:
            p.insert(idx + 1, make_run(value, use_red))
        # Insert after-text
        if after:
            p.insert(idx + 2, make_run(after, False))
    else:
        # Multiple occurrences in one run — replace all
        new_text = run.text.replace(code, value)
        run.text = new_text
        if use_red and value:
            run.font.color.rgb = RED


def consolidate_runs_for_code(paragraph, code):
    """If `code` is split across multiple <w:r> runs (which Word commonly does after
    edits — spell-check, autocorrect, copy-paste etc.), merge those runs into a single
    run that contains the whole code. The merged run keeps the formatting of the first
    affected run. No-op if the code is already in a single run, or not present.
    """
    runs = paragraph.runs
    if not runs:
        return

    full = "".join(r.text for r in runs)
    idx = full.find(code)
    if idx < 0:
        return  # Code not present
    end_idx = idx + len(code)

    pos = 0
    start_run_idx = None
    end_run_idx = None
    for i, r in enumerate(runs):
        run_end = pos + len(r.text)
        if start_run_idx is None and pos <= idx < run_end:
            start_run_idx = i
        if end_run_idx is None and pos < end_idx <= run_end:
            end_run_idx = i
            break
        pos = run_end

    if start_run_idx is None or end_run_idx is None:
        return
    if start_run_idx == end_run_idx:
        return  # Already in one run — fast path handles it

    # Merge text from runs (start+1 .. end) into the start run, then drop them.
    base = runs[start_run_idx]
    appended = "".join(runs[j].text for j in range(start_run_idx + 1, end_run_idx + 1))
    base.text = base.text + appended
    for j in range(end_run_idx, start_run_idx, -1):
        parent = runs[j]._r.getparent()
        if parent is not None:
            parent.remove(runs[j]._r)


def process_paragraph_text(paragraph, data, unmapped):
    """Replace all known codes in a paragraph. Leave unmapped codes untouched."""
    full = get_full_text(paragraph)
    if "{{" not in full:
        return

    # Find all codes in this paragraph
    codes_present = re.findall(r'\{\{[^}]+\}\}', full)

    for code in codes_present:
        if code in unmapped:
            continue  # Leave raw
        if code in data:
            value = data[code]
            # If Word split the code across multiple runs, merge them first so
            # `code in run.text` will succeed below.
            consolidate_runs_for_code(paragraph, code)
            # Work run by run
            for run in paragraph.runs:
                if code in run.text:
                    replace_code_in_run(run, code, value, use_red=True)
                    break  # one replacement per code per pass


def process_table(table, data, unmapped):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph_text(paragraph, data, unmapped)
            for nested_table in cell.tables:
                process_table(nested_table, data, unmapped)


def collect_all_paragraphs(doc):
    """Return flat list of (paragraph, parent_element, index) for body + tables."""
    items = []
    body = doc.element.body
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            items.append(Paragraph(child, doc))
        elif tag == 'tbl':
            from docx.table import Table
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        items.append(p)
    return items


def apply_conditional_deletions(doc, conditionals):
    """
    Walk through document body elements.
    When a start-tag paragraph is found and its condition is True,
    collect and remove all elements up to and including the end-tag paragraph.
    """
    body = doc.element.body
    elements = list(body)

    def get_para_text(el):
        return "".join(t.text or "" for t in el.iter(qn('w:t')))

    for start_tag, end_tag, condition_key in CONDITIONAL_PAIRS:
        should_delete = conditionals.get(condition_key, False)
        if not should_delete:
            # Still remove the marker tags themselves (they're not content)
            to_remove = []
            for el in list(body):
                txt = get_para_text(el)
                if start_tag in txt or end_tag in txt:
                    to_remove.append(el)
            for el in to_remove:
                body.remove(el)
            continue

        # Delete everything between (and including) start and end tags
        in_block = False
        to_remove = []
        for el in list(body):
            txt = get_para_text(el)
            if start_tag in txt:
                in_block = True
            if in_block:
                to_remove.append(el)
            if end_tag in txt and in_block:
                in_block = False
        for el in to_remove:
            try:
                body.remove(el)
            except ValueError:
                pass

    # Handle DeleteIfNoInsuranceAtAll (no end tag)
    # Remove the single marker tag paragraph regardless
    should_delete_no_ins = conditionals.get("DeleteIfNoInsuranceAtAll", False)
    to_remove = []
    in_block = False
    for el in list(body):
        txt = get_para_text(el)
        if NO_INSURANCE_SINGLE_TAG in txt:
            to_remove.append(el)  # always remove the tag itself
            if should_delete_no_ins:
                in_block = True
            continue
        if in_block:
            # Delete until we hit the next heading-level paragraph or end of section
            # Heuristic: stop at next paragraph that has bold text > 12pt or is a heading style
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag == 'p':
                style = el.find('.//' + qn('w:pStyle'))
                style_val = style.get(qn('w:val'), '') if style is not None else ''
                if 'Heading' in style_val or style_val.startswith('h'):
                    in_block = False
                    continue
            to_remove.append(el)
    for el in to_remove:
        try:
            body.remove(el)
        except ValueError:
            pass


def insert_scenario_content(doc, library, scenario_num):
    """
    Walk the SOA template body. For every paragraph that contains a scenario marker:
      - Non-matching scenarios' markers → always removed (clean up unused content).
      - Matching scenario's markers:
          * If library has real content for the marker → splice it in, remove marker.
          * If library is provided but the marker is empty / whitespace-only → strip marker cleanly.
          * If no library at all → leave the marker raw so the adviser knows where to fill in.

    Handles ScenarioNa..ScenarioNi, ScenarioNoptIn, ScenarioNAdviceLimitation.
    A no-op only if `scenario_num` is empty.
    """
    if not scenario_num:
        return
    target = str(scenario_num)
    have_library = bool(library)
    library = library or {}

    def _has_real_content(elements):
        for src in elements:
            t = src.tag.split('}')[-1] if '}' in src.tag else src.tag
            if t == 'tbl':
                return True
            if t == 'p' and _get_para_text(src).strip():
                return True
        return False

    body = doc.element.body
    # Snapshot to avoid mutation-during-iteration issues
    for el in list(body):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag != 'p':
            continue
        text = _get_para_text(el)
        m = SCENARIO_MARKER_RE.search(text)
        if not m:
            continue

        marker_scenario = m.group(1)   # '1'..'7'
        marker_full     = m.group()    # e.g. '{{Scenario1a}}'

        parent = el.getparent()
        if parent is None:
            continue

        if marker_scenario == target:
            content = library.get(marker_full, [])
            if _has_real_content(content):
                # Library has real content for this marker — splice it in, remove marker
                idx = list(parent).index(el)
                for i, src_el in enumerate(content):
                    parent.insert(idx + 1 + i, copy.deepcopy(src_el))
                parent.remove(el)
            elif have_library:
                # Library uploaded but this marker is empty (e.g. Scenario 5 "NA")
                # — strip the marker so the output is clean
                parent.remove(el)
            # else: no library uploaded → leave marker raw so adviser fills it in
        else:
            # Non-matching scenario marker — always strip
            parent.remove(el)


def process_soa(template_bytes, data, conditionals, scenario_library=None, scenario_num=""):
    """Main processor — returns completed docx as bytes."""
    import gc

    # Load document then immediately free the raw bytes from memory
    buf = io.BytesIO(template_bytes)
    doc = Document(buf)
    del template_bytes, buf
    gc.collect()

    # Build a per-call unmapped set: any code that ALSO has a non-empty value in
    # the data dict (e.g. a goal injected from the KYC note) is removed from the
    # unmapped set so it gets replaced rather than left raw.
    runtime_unmapped = UNMAPPED_CODES - {k for k, v in data.items() if v}

    # Step 1: Splice scenario content from the bundled library FIRST.
    # This brings in any conditional markers ({{DeleteIfNoScopedTrauma}} etc.) that
    # live inside scenario sub-blocks so the conditional pass below can see and
    # process them. Runs even without a library — unused scenarios' markers are
    # stripped so they don't pollute the output.
    if scenario_num:
        insert_scenario_content(doc, scenario_library, scenario_num)

    # Step 1b: Apply ALL conditional block deletions across the fully-spliced document.
    # Process is: add everything in first (Step 1), then prune unwanted content last.
    # This means every DeleteIf* marker is processed regardless of whether it lived
    # at the top level of the template or was injected via scenario content.
    apply_conditional_deletions(doc, conditionals)

    # Step 2: Replace codes in body paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph_text(paragraph, data, runtime_unmapped)

    # Step 3: Replace codes in tables
    for table in doc.tables:
        process_table(table, data, runtime_unmapped)

    # Step 4: Replace codes in headers and footers
    for section in doc.sections:
        for hdr in [section.header, section.footer,
                    section.even_page_header, section.even_page_footer,
                    section.first_page_header, section.first_page_footer]:
            if hdr:
                for paragraph in hdr.paragraphs:
                    process_paragraph_text(paragraph, data, runtime_unmapped)
                for table in hdr.tables:
                    process_table(table, data, runtime_unmapped)

    out = io.BytesIO()
    doc.save(out)
    del doc
    gc.collect()
    out.seek(0)
    return out


# ─────────────────────────────────────────────
# LOGIN PAGE HTML
# ─────────────────────────────────────────────

LOGIN_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>SOA Agent — Brightday Australia</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Ubuntu:wght@300;400;500;700&display=swap" rel="stylesheet">
<link href="https://api.fontshare.com/v2/css?f[]=general-sans@300,400,500,600,700&display=swap" rel="stylesheet">
<style>
  :root {
    --navy:        #123559;
    --navy-deep:   #0C243D;
    --raspberry:   #F50D74;
    --raspberry-light: #FF4593;
    --bg:          #FFFFFF;
    --surface:     #FFFFFF;
    --surface-soft:#F7F8FA;
    --text:        #0E2640;
    --text-muted:  #5C6B7E;
    --text-soft:   #8A95A4;
    --border:      #E5E9EF;
    --border-strong:#CBD2DB;
    --red:         #DA2929;
    --red-soft:    #FCEEEE;
  }
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  html, body { height: 100%; }
  body {
    background: var(--bg);
    color: var(--text);
    font-family: 'General Sans', 'Ubuntu', system-ui, sans-serif;
    font-weight: 400;
    min-height: 100vh;
    display: flex; flex-direction: column;
  }
  body::before {
    content: ''; position: fixed; top: -180px; right: -180px;
    width: 520px; height: 520px;
    background: radial-gradient(circle, rgba(245,13,116,0.10) 0%, transparent 70%);
    pointer-events: none; z-index: 0;
  }
  body::after {
    content: ''; position: fixed; bottom: -200px; left: -200px;
    width: 520px; height: 520px;
    background: radial-gradient(circle, rgba(18,53,89,0.06) 0%, transparent 70%);
    pointer-events: none; z-index: 0;
  }
  .site-header {
    background: var(--navy);
    border-bottom: 3px solid var(--raspberry);
    z-index: 1; position: relative;
  }
  .header-inner {
    max-width: 1080px; margin: 0 auto;
    padding: 18px 32px;
    display: flex; align-items: center; gap: 14px;
  }
  .logo-mark {
    width: 38px; height: 38px;
    display: flex; align-items: center; justify-content: center; flex-shrink: 0;
  }
  .logo-mark svg { width: 100%; height: 100%; display: block; }
  .header-text { display: flex; flex-direction: column; line-height: 1; }
  .header-text .brand {
    font-family: 'Ubuntu', sans-serif; font-size: 22px; font-weight: 700;
    color: #FFFFFF; letter-spacing: -0.5px;
  }
  .header-text .sub {
    font-family: 'General Sans', sans-serif;
    font-size: 10px; color: var(--raspberry-light); letter-spacing: 2.2px; font-weight: 500;
    text-transform: uppercase; margin-top: 5px;
  }

  .login-wrap {
    flex: 1;
    display: flex; align-items: center; justify-content: center;
    padding: 48px 24px;
    position: relative; z-index: 1;
  }
  .card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 44px 40px;
    width: 100%; max-width: 420px;
    animation: fadeUp 0.6s ease both;
    box-shadow: 0 12px 40px rgba(14,38,64,0.08);
  }
  h1 {
    font-family: 'Ubuntu', sans-serif; font-size: 28px; font-weight: 700;
    color: var(--navy);
    line-height: 1.1; margin-bottom: 8px; letter-spacing: -0.6px;
  }
  h1 em { font-style: normal; color: var(--raspberry); }
  .subtitle {
    font-family: 'General Sans', sans-serif;
    font-size: 13px; color: var(--text-muted); margin-bottom: 32px; line-height: 1.6;
  }
  .field { margin-bottom: 18px; }
  .field label {
    font-family: 'General Sans', sans-serif;
    display: block; font-size: 10px; font-weight: 600;
    letter-spacing: 2px; text-transform: uppercase;
    color: var(--text-muted); margin-bottom: 6px;
  }
  .field input {
    width: 100%;
    background: var(--surface);
    border: 1px solid var(--border);
    color: var(--text);
    font-family: 'General Sans', sans-serif; font-size: 13px;
    padding: 12px 14px; outline: none;
    border-radius: 8px;
    transition: border-color 0.2s, box-shadow 0.2s;
  }
  .field input:hover { border-color: var(--border-strong); }
  .field input:focus { border-color: var(--raspberry); box-shadow: 0 0 0 3px rgba(245,13,116,0.10); }
  .btn-login {
    width: 100%;
    background: var(--raspberry); color: #FFFFFF; border: none;
    padding: 14px;
    font-family: 'Ubuntu', sans-serif; font-size: 11px;
    font-weight: 700; letter-spacing: 2px; text-transform: uppercase;
    cursor: pointer; margin-top: 12px;
    border-radius: 999px;
    box-shadow: 0 4px 14px rgba(245,13,116,0.30);
    transition: background 0.2s, box-shadow 0.2s;
  }
  .btn-login:hover { background: var(--raspberry-light); box-shadow: 0 6px 18px rgba(245,13,116,0.36); }
  .error {
    background: var(--red-soft);
    border: 1px solid rgba(218,41,41,0.30);
    color: var(--red);
    font-size: 12px; padding: 10px 14px; margin-bottom: 18px;
    border-radius: 8px;
    font-family: 'General Sans', sans-serif;
  }
  .footer-note {
    font-family: 'General Sans', sans-serif;
    font-size: 10px; color: var(--text-soft);
    margin-top: 28px; text-align: center; letter-spacing: 0.5px;
  }
  .site-footer {
    background: var(--navy-deep);
    color: rgba(245,247,250,0.72);
    padding: 18px 32px;
    text-align: center;
    z-index: 1; position: relative;
  }
  .site-footer .footer-brand {
    font-family: 'Ubuntu', sans-serif;
    font-size: 13px; font-weight: 700; color: #F5F7FA;
    margin-right: 12px;
  }
  .site-footer .footer-brand span { color: var(--raspberry-light); }
  .site-footer .footer-copy {
    font-family: 'General Sans', sans-serif;
    font-size: 10px; opacity: 0.85; letter-spacing: 0.5px;
  }
  @keyframes fadeUp { from { opacity: 0; transform: translateY(16px); } to { opacity: 1; transform: translateY(0); } }
</style>
</head>
<body>

<header class="site-header">
  <div class="header-inner">
    <div class="logo-mark" aria-label="Brightday logomark">
      <svg viewBox="0 0 100 100" xmlns="http://www.w3.org/2000/svg" role="img" aria-hidden="true">
        <g transform="translate(50,50)">
          <g fill="#F50D74" fill-opacity="0.78">
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(0)"/>
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(45)"/>
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(90)"/>
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(135)"/>
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(180)"/>
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(225)"/>
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(270)"/>
            <rect x="-13" y="-34" width="26" height="26" rx="7" transform="rotate(315)"/>
          </g>
        </g>
      </svg>
    </div>
    <div class="header-text">
      <span class="brand">brightday</span>
      <span class="sub">SOA Completion Agent</span>
    </div>
  </div>
</header>

<main class="login-wrap">
  <div class="card">
    <h1>Sign <em>in</em></h1>
    <p class="subtitle">Internal access only. Enter your credentials to continue.</p>
    {% if error %}
    <div class="error">{{ error }}</div>
    {% endif %}
    <form method="POST" action="/login">
      <div class="field">
        <label>Username</label>
        <input type="text" name="username" autocomplete="username" autofocus required>
      </div>
      <div class="field">
        <label>Password</label>
        <input type="password" name="password" autocomplete="current-password" required>
      </div>
      <button class="btn-login" type="submit">Sign In →</button>
    </form>
    <p class="footer-note">Brightday Australia · ABN 45 674 252 905</p>
  </div>
</main>

<footer class="site-footer">
  <span class="footer-brand">brightday <span>Australia</span></span>
  <span class="footer-copy">ABN 45 674 252 905 · 260 Spencer Street, Melbourne · Internal Use Only</span>
</footer>

</body>
</html>"""


# ─────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        if check_password(username, password):
            session["authenticated"] = True
            session["username"] = username.lower()
            return redirect(url_for("tool"))
        else:
            error = "Invalid username or password."
    return render_template_string(LOGIN_HTML, error=error)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/")
def tool():
    if not logged_in():
        return redirect(url_for("login"))
    html_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "index.html")
    with open(html_path, "r") as f:
        return f.read()


@app.route("/api/extract", methods=["POST"])
def api_extract():
    """
    Step 1 of the wizard: parse the KYC note (and optionally peek at the FF) to
    return suggested UI defaults for Step 2. The FF and SOA template are NOT
    persisted here — the browser keeps the File objects and re-sends them on /process.

    Request: multipart form with optional 'kyc_note' file (.docx).
    Response JSON:
        {
          "kyc": { ... read_kyc_note output ... } or null,
          "suggested": {
            "risk_profile": "...",
            "no_insurance":   bool,
            "no_salsac":      bool,
            "no_trauma":      bool,
            "insurance_only": bool,
          },
          "conflicts": []     # placeholder for future FF↔KYC conflict detection
        }
    """
    if not logged_in():
        return jsonify({"error": "Not authenticated"}), 401
    try:
        kyc_data = None
        if "kyc_note" in request.files and request.files["kyc_note"].filename:
            kyc_bytes = request.files["kyc_note"].read()
            kyc_data = read_kyc_note(kyc_bytes)
            del kyc_bytes

        suggested = {
            "risk_profile":   "",
            "no_insurance":   False,
            "no_salsac":      False,
            "no_trauma":      False,
            "insurance_only": False,
        }
        if kyc_data:
            suggested["risk_profile"] = kyc_data.get("risk_profile", "")
            scope = kyc_data.get("scope", {})
            # Map scope tags to UI-checkbox defaults. 'in' = leave unticked, 'out' = pre-tick.
            suggested["no_insurance"] = (scope.get("insurance") == "out")
            suggested["no_salsac"]    = (scope.get("salary_sacrifice") == "out")
            # No KYC scope tag exists for trauma or insurance-only; adviser sets manually.

        return jsonify({
            "kyc":        kyc_data,
            "suggested":  suggested,
            "conflicts":  [],   # Phase 3b: FF vs KYC conflict detection
        })
    except Exception as e:
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500


@app.route("/process", methods=["POST"])
def process():
    if not logged_in():
        return jsonify({"error": "Not authenticated"}), 401
    try:
        if "fact_finder" not in request.files:
            return jsonify({"error": "Missing Fact Finder file"}), 400
        if "soa_template" not in request.files:
            return jsonify({"error": "Missing SOA Template file"}), 400

        risk_profile   = request.form.get("risk_profile", "").strip()
        scenario       = request.form.get("scenario", "").strip()
        rollover_type  = request.form.get("rollover_type", "").strip()   # "Full" / "Partial" / ""
        no_insurance   = request.form.get("no_insurance", "false").lower() == "true"
        no_trauma      = request.form.get("no_trauma", "false").lower() == "true"
        no_salsac      = request.form.get("no_salsac", "false").lower() == "true"
        insurance_only = request.form.get("insurance_only", "false").lower() == "true"

        # Optional goal text from KYC note (or adviser-edited at the review step).
        # Empty values mean "leave the placeholder raw for adviser to complete".
        goal_overrides = {
            "super":            request.form.get("goal_super", "").strip(),
            "insurance":        request.form.get("goal_insurance", "").strip(),
            "salary_sacrifice": request.form.get("goal_salary_sacrifice", "").strip(),
            "estate_planning":  request.form.get("goal_estate_planning", "").strip(),
            "retirement":       request.form.get("goal_retirement", "").strip(),
        }

        if not risk_profile:
            return jsonify({"error": "Risk profile must be selected"}), 400
        if not scenario:
            return jsonify({"error": "Scenario must be selected"}), 400
        if scenario not in {"1", "2", "3", "4", "5", "6", "7"}:
            return jsonify({"error": "Scenario must be 1-7"}), 400

        ff_bytes       = request.files["fact_finder"].read()
        template_bytes = request.files["soa_template"].read()
        # OFA template is optional. If supplied (non-empty filename), generate the OFA too
        # and bundle both files into a single zip for download.
        ofa_bytes = None
        if "ofa_template" in request.files and request.files["ofa_template"].filename:
            ofa_bytes = request.files["ofa_template"].read()

        # Scenario library — starts from the bundled default (scenario_library.docx that
        # ships alongside app.py). If the adviser uploads a different one via the UI,
        # that overrides for this request only. The default is what advisers will use
        # most of the time — no need to re-upload the firm-wide training doc per client.
        scenario_library = DEFAULT_SCENARIO_LIBRARY
        if "scenario_library" in request.files and request.files["scenario_library"].filename:
            sl_bytes = request.files["scenario_library"].read()
            scenario_library = read_scenario_library(sl_bytes)
            del sl_bytes

        # Read fact finder then free its bytes
        data, conditionals = read_fact_finder(
            ff_bytes, risk_profile, no_insurance,
            no_trauma_flag=no_trauma,
            no_salsac_flag=no_salsac,
            insurance_only_flag=insurance_only,
            scenario=scenario,
            goal_overrides=goal_overrides,
            rollover_type=rollover_type,
        )
        del ff_bytes

        client_name = data.get("{{ClientFullName}}", "Client")
        client_slug = client_name.replace(" ", "_") or "Client"
        today_str   = date.today().strftime("%Y%m%d")
        soa_name    = f"SOA_Draft_{client_slug}_{today_str}.docx"
        ofa_name    = f"OFA_Draft_{client_slug}_{today_str}.docx"
        DOCX_MIME   = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Generate the SOA (template_bytes is freed inside process_soa)
        soa_out = process_soa(template_bytes, data, conditionals,
                              scenario_library=scenario_library, scenario_num=scenario)

        if ofa_bytes is None:
            # Single-file response — preserves the legacy behavior.
            return send_file(soa_out, as_attachment=True, download_name=soa_name, mimetype=DOCX_MIME)

        # Both templates supplied: process the OFA as well, then zip both.
        # OFA template doesn't have scenario markers — scenario library is irrelevant there.
        ofa_out = process_soa(ofa_bytes, data, conditionals)
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(soa_name, soa_out.getvalue())
            zf.writestr(ofa_name, ofa_out.getvalue())
        zip_buf.seek(0)
        bundle_name = f"Documents_{client_slug}_{today_str}.zip"
        return send_file(zip_buf, as_attachment=True, download_name=bundle_name, mimetype="application/zip")

    except KeyError as e:
        return jsonify({"error": f"Fact Finder tab not found or unexpected structure: {e}"}), 400
    except Exception as e:
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=False, host="0.0.0.0", port=port)
